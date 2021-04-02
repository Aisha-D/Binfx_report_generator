from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import argparse
import datetime
import os 

def parse_args():
    parser = argparse.ArgumentParser()

    parser.add_argument(
        '-p', '--data_path',
        help='path to the data folder',
        required=True
        )

    args = parser.parse_args()

    return args

def Format_date(folder_date):
    date_y = folder_date[-4:-2] #subract fourth to second last char
    date_m = folder_date[-2:] #substracts last two char
    date = date_y + " " + date_m
    folder_date = date_y + date_m
    date_formatted = datetime.datetime.strptime(date, "%y %m")
    date = date_formatted.strftime("%b %Y")
    return date, folder_date

def Open_document(current_path):
    document = Document(current_path + '/templates/Report_template.docx')
    return document

## Add date under header
def Add_date(date, document):
    p = document.add_paragraph(date)
    p.alignment = 1
    p.add_run().size = Pt(14)
    document.add_paragraph(' ')

## Add Agents to user ratio
def Add_agent2user(document):
    table = document.add_table(rows=1, cols=2)
    table.style = 'Grid Table 4 Accent 5'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Agents'
    hdr_cells[1].text = 'Users'
    row_cells = table.add_row().cells
    row_cells[0].text = '11'
    row_cells[1].text = '51'
    document.add_paragraph(' ')

## Add helpdesk header
def Add_HelpdeskSection(document, path):
    # document.add_heading('Helpdesk', level=1)
    document.add_paragraph(' ')
    ### Helpdesk workload
    # document.add_heading('Workload')
    document.add_paragraph(
        'Since August 2020, we moved to the EMEE helpdesk solve to any bioinformatics related issues or queries. '
        'Steadily, more staff are raising issues via the helpdesk.'
        )
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/Workload.png', width=Inches(6))
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    document.add_paragraph(
        "This month's workload can be categorised into types such as: "
        "Submit a request or incident, Ask a question, Emailed request, "
        "Standard Reanalysis and  Urgent Reanalysis. From month to month, "
        "certain types of tickets may be requested more than others."
        )
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/Months_workload.png', width=Inches(6))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    ### Number of tickets created vs completed this month
    document.add_paragraph(
        'The helpdesk is designed to categorise different tickets into groups to help us understand what common issues are .'
        'We can see reanalysis, both standard and urgent, are highly requested and we meet respond back quickly within the month if there are no issues.'
        'Cases where there are more completed tickets than raised, the closed tickets are those raised in the last month.'
        )
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/Number_of_Tickets_Created_vs_Completed.png', width=Inches(6))
    last_paragraph = document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(' ')
    #### SLAs
    # document.add_heading('SLAs')
    document.add_paragraph(
        'We set time goals with Service Level Agreements (SLAs) to '
        'help drive better quality of service. Some issues are set '
        'different times. The times for out tickets '
        'work on a 24/7 calendar.'
        )
    document.add_paragraph(' ')

    tickets = (
    ('Standard Reanalysis', '72', '336'),
    ('Urgent Reanalysis',  '72', '120'),
    ('Submit a request or incident', '72', '336'),
    ('Ask a question', '72', '336'),
    ('Emailed request', '72', '336'),
    ('Inform us', '72', '336')
    )
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticket_types'
    hdr_cells[1].text = 'Time_to_first_response'
    hdr_cells[2].text = 'Time_to_resolution'
    for Ticket_types, Time_to_first_response, Time_to_resolution in tickets:
        row_cells = table.add_row().cells
        row_cells[0].text = Ticket_types
        row_cells[1].text = Time_to_first_response
        row_cells[2].text = Time_to_resolution
    hdr_cells[0].text = ' '
    hdr_cells[1].text = 'Time to first response (hrs)'
    hdr_cells[2].text = 'Time to resolution (hrs)'
    
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    document.add_paragraph(
        'We aim to meet within target hours of our tickets, '
        'however, some SLAs may be breached due to the task being '
        'difficult hence taking longer to resolve.'
        " Most SLAs breached are 'Submit a request or incident' "
        ' as they tend to require more time to resolve. '
        )
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/resolution_SLAs.png', width=Inches(6))
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/response_SLAs.png', width=Inches(6))
    document.add_paragraph(' ')
    document.add_paragraph(' ')

    #### Current statuses
    # document.add_heading('Current Statuses')
    document.add_paragraph(' ')
    document.add_paragraph(' ')
    document.add_paragraph(
        'Our helpdesk is busy and to track progress of our work, '
        'tickets are places into four categories: To do, In Progress, '
        'On hold and Waiting for response. Here is a snapshot of '
        "our current ticket's statuses."
        )
    document.add_paragraph(' ')
    my_image = document.add_picture(path+'/Current_Statuses.png', width=Inches(6))


def Close_document(document, path, date):
    document.save(path + '/Report_' + date + '.docx')
    print("Updated the report and saved in: " + path + '/Report_' + date + '.docx')


def main():

    current_path = os.getcwd()
    #current_path = current_path[:-5]

    args = parse_args()

    date, folder_date = Format_date(args.data_path)

    path = current_path + '/data/' + folder_date

    print(date)

    print(folder_date)

    print(path)

    document = Open_document(current_path)
    
    Add_date(date, document)

    Add_agent2user(document)

    Add_HelpdeskSection(document, path)
    
    date = date.replace(" ", "_")

    Close_document(document, path, date)


if __name__ == "__main__":

    main()
